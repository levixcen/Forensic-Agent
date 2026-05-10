import os
import sys
import io
import json
from flask import Flask, render_template, request, jsonify, send_file, Response, stream_with_context
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv()

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from agent.main import run_agent, call_azure_openai

app = Flask(__name__)
app.secret_key = os.urandom(24)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/analyze', methods=['POST'])
def analyze():
    if 'logfile' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['logfile']
    if not file.filename:
        return jsonify({'error': 'Empty filename'}), 400

    filepath = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(filepath)

    def generate():
        try:
            report = run_agent(filepath)

            # ----------------------------------------------------------------
            # Stream word-by-word so the browser updates progressively.
            # Each SSE frame is  data: {"token": "word "}
            # A final frame     data: {"done": true}  signals completion.
            #
            # Fixes:
            #   - Every yield ends with \n\n  (required by SSE spec)
            #   - We never swallow the done frame even on empty reports
            #   - Errors are forwarded as  data: {"error": "..."}
            # ----------------------------------------------------------------
            words = report.split(' ')
            for word in words:
                payload = json.dumps({'token': word + ' '})
                yield f"data: {payload}\n\n"

            yield 'data: {"done": true}\n\n'

        except Exception as e:
            error_payload = json.dumps({'error': str(e)})
            yield f"data: {error_payload}\n\n"
            yield 'data: {"done": true}\n\n'   # always close the stream

    response = Response(
        stream_with_context(generate()),
        mimetype='text/event-stream'
    )
    # Prevent proxy / nginx buffering from eating the stream
    response.headers['X-Accel-Buffering'] = 'no'
    response.headers['Cache-Control'] = 'no-cache'
    return response


@app.route('/chat', methods=['POST'])
def chat():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data received'}), 400

    question = data.get('question', '').strip()
    report_context = data.get('report_context', '').strip()

    if not question:
        return jsonify({'error': 'No question provided'}), 400
    if not report_context:
        return jsonify({'error': 'No report context provided'}), 400

    prompt = f"""
You are a senior digital forensic analyst. A forensic investigation report
has already been generated and is shown below. The analyst is asking a
follow-up question about this specific incident.

FORENSIC REPORT:
{report_context}

ANALYST QUESTION:
{question}

Answer the question based strictly on the evidence in the report above.
Be concise and specific. Do not use markdown formatting, no asterisks,
no hashtags. Use plain text only. If the answer requires a list,
use numbers like 1. 2. 3.
If the question cannot be answered from the report, say so clearly.
"""
    try:
        response = call_azure_openai(prompt)
        return jsonify({'response': response})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/download', methods=['POST'])
def download():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data received'}), 400

    report_text = data.get('report', '')
    if not report_text:
        return jsonify({'error': 'No report text provided'}), 400

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = doc.add_heading('LogIQ Forensic Investigation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1A, 0x5F, 0xA5)
    doc.add_paragraph()

    sections = [
        'WHAT HAPPENED AND WHEN',
        'WHAT WAS THE ROOT CAUSE',
        'ATTACK TIMELINE',
        'MITRE ATT&CK MAPPING',
        'WHAT WAS AND REMAINS TO BE DONE',
        'LESSONS LEARNED',
        'RECOMMENDED NEXT STEPS',
        'ANALYST NOTES',
    ]
    meta_fields = ['CLASSIFICATION', 'REPORT ID', 'DATE', 'ANALYST']

    lines = report_text.strip().split('\n')
    body_lines = []
    in_meta = True
    meta_pairs = []

    for line in lines:
        trimmed = line.strip()
        if not trimmed:
            if not in_meta:
                body_lines.append('')
            continue
        is_meta = any(trimmed.upper().startswith(f + ':') for f in meta_fields)
        if is_meta and in_meta:
            colon = trimmed.index(':')
            meta_pairs.append((trimmed[:colon].strip(), trimmed[colon + 1:].strip()))
        else:
            in_meta = False
            body_lines.append(trimmed)

    if meta_pairs:
        table = doc.add_table(rows=len(meta_pairs), cols=2)
        table.style = 'Table Grid'
        for i, (key, val) in enumerate(meta_pairs):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = val
            runs = table.rows[i].cells[0].paragraphs[0].runs
            if runs:
                runs[0].bold = True
                runs[0].font.color.rgb = RGBColor(0x37, 0x8A, 0xDD)
        doc.add_paragraph()

    body_text = '\n'.join(body_lines)
    for section in sections:
        body_text = body_text.replace(section, f'|||SECTION|||{section}')

    parts = body_text.split('|||SECTION|||')
    for part in parts:
        if not part.strip():
            continue
        matched = next((s for s in sections if part.startswith(s)), None)
        if matched:
            first_newline = part.index('\n') if '\n' in part else len(part)
            content = part[first_newline:].strip()

            heading = doc.add_heading(matched, level=1)
            if heading.runs:
                heading.runs[0].font.color.rgb = RGBColor(0x37, 0x8A, 0xDD)
                heading.runs[0].font.size = Pt(12)

            if matched == 'MITRE ATT&CK MAPPING' and content:
                mitre_lines = [
                    l.strip() for l in content.split('\n')
                    if l.strip() and '|' in l and not re.match(r'^\|[-: |]+\|$', l.strip())
                ]
                if mitre_lines:
                    t = doc.add_table(rows=1 + len(mitre_lines), cols=4)
                    t.style = 'Table Grid'
                    headers = ['TACTIC', 'TECHNIQUE ID', 'TECHNIQUE NAME', 'SEVERITY']
                    _write_table_header(t, headers)
                    for i, line in enumerate(mitre_lines):
                        cells = [c.strip() for c in line.strip('|').split('|')]
                        for j in range(4):
                            t.rows[i + 1].cells[j].text = cells[j] if j < len(cells) else ''

            elif matched == 'ATTACK TIMELINE' and content:
                timeline_lines = [
                    l.strip() for l in content.split('\n')
                    if l.strip() and '|' in l and not re.match(r'^\|[-: |]+\|$', l.strip())
                ]
                if timeline_lines:
                    t = doc.add_table(rows=1 + len(timeline_lines), cols=3)
                    t.style = 'Table Grid'
                    headers = ['TIMESTAMP', 'EVENT ID', 'DESCRIPTION']
                    _write_table_header(t, headers)
                    for i, line in enumerate(timeline_lines):
                        cells = [c.strip() for c in line.strip('|').split('|')]
                        for j in range(3):
                            t.rows[i + 1].cells[j].text = cells[j] if j < len(cells) else ''

            elif matched in (
                'WHAT HAPPENED AND WHEN', 'WHAT WAS THE ROOT CAUSE',
                'WHAT WAS AND REMAINS TO BE DONE', 'LESSONS LEARNED'
            ) and content:
                kv_lines = [l.strip() for l in content.split('\n') if l.strip()]
                kv_pairs = []
                for l in kv_lines:
                    colon = l.find(':')
                    if 0 < colon < 40:
                        kv_pairs.append((l[:colon].strip(), l[colon + 1:].strip()))
                    else:
                        kv_pairs.append(('', l))
                if kv_pairs:
                    t = doc.add_table(rows=len(kv_pairs), cols=2)
                    t.style = 'Table Grid'
                    for i, (key, val) in enumerate(kv_pairs):
                        t.rows[i].cells[0].text = key
                        t.rows[i].cells[1].text = val
                        if key:
                            runs = t.rows[i].cells[0].paragraphs[0].runs
                            if runs:
                                runs[0].bold = True
                                runs[0].font.color.rgb = RGBColor(0x37, 0x8A, 0xDD)

            elif content:
                for line in content.split('\n'):
                    l = line.strip()
                    if not l:
                        continue
                    p = doc.add_paragraph(l)
                    p.style = 'List Number' if (l[0].isdigit() and '. ' in l[:4]) else 'Normal'

            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name='LogIQ_Forensic_Report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _write_table_header(table, headers):
    """Write a bold white-on-blue header row."""
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        runs = cell.paragraphs[0].runs
        if runs:
            runs[0].bold = True
            runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '185FA5')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)


import re  # needed inside download() for the markdown-separator filter


if __name__ == '__main__':
    app.run(debug=True)